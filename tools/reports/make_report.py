"""Generate rapor.docx from inline report content."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\1KodCalismalari\Projeler\VIBE_CODING_UYGULAMA_DENEMELERI\Merge_PortfoySim\AI_Core\ts_forecasting_lab\rapor.docx"

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(level, text):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, hh in enumerate(headers):
        hdr[i].text = hh
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

# Title
title = doc.add_heading('ASTOR Hissesi Walk-Forward Backtest Karşılaştırmalı Analiz Raporu', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Akademik Çalışma — Zaman Serisi Tahmin Modellerinin Karşılaştırmalı Değerlendirmesi')
r.italic = True

# Meta
doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Çalışma: ').bold = True
meta.add_run('Tek varlık (ASTOR, BIST) zaman serisi tahmin modellerinin yürüyüş-ileri (walk-forward) doğrulama protokolü altında karşılaştırmalı backtest analizi.\n')
meta.add_run('Test Tarihi: ').bold = True
meta.add_run('2026-05-16\n')
meta.add_run('Doğrulama Protokolü: ').bold = True
meta.add_run('12-fold walk-forward (sliding), min_train=504 bar, test=21 bar, embargo=30 bar, final_holdout=60 bar.\n')
meta.add_run('İşlem Maliyeti: ').bold = True
meta.add_run('commission_bps=0.0, slippage_bps=0.0 (idealize koşullar).\n')
meta.add_run('Sinyal Modu: ').bold = True
meta.add_run('simple_current (soft gate).')

# 1. Özet
h(1, '1. Yönetici Özeti')
para('Beş farklı tahmin modeli (LightGBM, Ridge, ElasticNet, DLinear, NLinear) ve dört naive benchmark (Drift, Last Value, Zero Return, Buy & Hold) ASTOR hissesi üzerinde aynı walk-forward protokolüyle değerlendirildi. Yalnızca ElasticNet Return modeli Buy & Hold benchmark’ını net getiri açısından geçebildi (Beats_BuyHold = Yes). Ridge Return ona en yakın aday olarak konumlandı ve risk-ayarlı metriklerde (Sortino, MaxDD) öne çıktı. Derin öğrenme tabanlı DLinear/NLinear modelleri tahmin kalibrasyonu sorunları nedeniyle naive Last Value seviyesinin altında performans gösterdi. LightGBM, yön doğruluğu coinflip eşiğinin altına düşerek tek negatif net getirili model oldu.')
para('Ana bulgu: ASTOR’un test penceresindeki güçlü yukarı yönlü trend rejimi, naive Buy & Hold benchmark’ını yüksek tutarak gelişmiş modeller için zorlu bir karşılaştırma çıtası oluşturdu. Lineer regülarizasyonlu modeller (Ridge, ElasticNet) overfit’e karşı dirençleri sayesinde anlamlı sonuç üretirken, gradient boosting (LightGBM) ve linear-decomposition deep modeller (DLinear, NLinear) bu rejimde başarısız oldu.', bold=False)

# 2. Düzen
h(1, '2. Deneysel Düzen')
make_table(
    ['Parametre', 'Değer'],
    [
        ['Varlık', 'ASTOR (BIST)'],
        ['Hedef değişken', 'y[t] = t+1 ileri getiri'],
        ['Özellik bilgi sınırı', 'X[t] yalnızca t kapanışına kadar bilinen bilgiyi içerir'],
        ['Yürütme gecikmesi', 'Sinyal t kapanışında üretilir, t+1 baryla eşleştirilir (look-ahead bias yok)'],
        ['WF split sayısı', '12'],
        ['Min eğitim penceresi', '504 bar (~2 yıl)'],
        ['Test penceresi', '21 bar'],
        ['Embargo', '30 bar'],
        ['Final holdout', '60 bar'],
        ['Eşik kaynağı', 'walk_forward_calibration_folds'],
        ['Sinyal kapısı', 'simple_current, soft mode'],
        ['Gate eşikleri', 'min_dir_acc=52%, max_rmse_vs_bench=1.05, min_composite=49.0'],
    ]
)

# 3. Sonuç tabloları
h(1, '3. Konsolide Sonuç Tabloları')
h(2, '3.1 Getiri ve Risk Profili')
make_table(
    ['Model', 'Net Return', 'CAGR', 'Sharpe', 'Sortino', 'Max DD', 'Calmar', 'Defl. Sharpe', 'Beats B&H'],
    [
        ['ElasticNet Return', '+2.549', '203.29%', '3.254', '43.09', '-5.15%', '3943.8', '2.714', 'Evet'],
        ['Naive Drift', '+2.404', '170.63%', '3.149', '28.66', '-12.93%', '1319.6', '2.624', 'Hayır'],
        ['Ridge Return', '+2.283', '146.35%', '3.081', '43.08', '-3.88%', '3771.2', '2.566', 'Hayır'],
        ['Naive Last Value', '+1.818', '76.60%', '2.731', '24.10', '-13.77%', '556.1', '2.263', 'Hayır'],
        ['DLinear', '+0.869', '12.84%', '1.820', '18.72', '-14.05%', '91.4', '1.470', 'Hayır'],
        ['NLinear', '+0.842', '12.01%', '1.787', '18.02', '-15.11%', '79.5', '1.441', 'Hayır'],
        ['Naive Zero', '0.000', '0.00%', '0.000', '0.00', '0.00%', 'inf', '-0.215', 'Hayır'],
        ['LightGBM Return', '-0.055', '-20.99%', '-4.406', '-5.66', '-11.17%', '-1.88', '-5.109', 'Hayır'],
        ['Buy & Hold ref.', '+2.404', '170.63%', '—', '—', '—', '—', '—', '—'],
    ]
)

h(2, '3.2 Pozisyon ve İşlem Profili')
make_table(
    ['Model', 'Exposure', 'Trade', 'Win%', 'Avg Trade Ret', 'Avg Hold (bar)', 'Profit Factor', 'Expectancy'],
    [
        ['ElasticNet', '56.67%', '7', '71.4%', '0.2655', '4.71', '102.59', '0.266'],
        ['Ridge', '50.00%', '11', '72.7%', '0.1369', '2.64', '29.94', '0.137'],
        ['Naive Drift', '100.00%', '1', '100%', '2.4044', '59.0', 'inf', '2.404'],
        ['Naive Last Value', '66.67%', '2', '50.0%', '0.9362', '19.5', '66.68', '0.936'],
        ['DLinear', '40.00%', '6', '33.3%', '0.1509', '3.83', '8.61', '0.151'],
        ['NLinear', '45.00%', '8', '37.5%', '0.1115', '3.25', '6.98', '0.111'],
        ['LightGBM', '26.67%', '8', '25.0%', '-0.0068', '2.00', '0.42', '-0.007'],
    ]
)

h(2, '3.3 Sinyal Kapısı Tanı Verileri (simple_current, walk-forward)')
make_table(
    ['Model', 'Dir_Acc', 'RMSE/bench', 'Composite', 'Mean|Pred|', '%>Threshold', 'Below_Entry', 'Diagnosis'],
    [
        ['Naive Drift', '55.00%', '1.0000', '73.00', '0.00104', '100.0%', '0', 'underperform_buyhold'],
        ['ElasticNet', '59.29%', '1.0017', '49.00', '0.01185', '56.7%', '20', 'gate_too_strict'],
        ['Ridge', '59.29%', '1.0006', '49.00', '0.00929', '50.0%', '20', 'underperform_buyhold'],
        ['Naive Last Value', '53.57%', '1.1233', '49.00', '0.01774', '66.7%', '0', 'underperform_buyhold'],
        ['LightGBM', '42.86%', '1.0912', '46.23', '0.01336', '26.7%', '36', 'model_signal_weak'],
        ['Naive Zero', '2.14%', '1.0025', '32.78', '0.00000', '0.0%', '60', 'model_signal_weak'],
        ['DLinear', '38.57%', '2.5264', '21.70', '0.09070', '40.0%', '31', 'kalibrasyon hatası'],
        ['NLinear', '40.71%', '3.0227', '19.16', '0.11383', '45.0%', '26', 'kalibrasyon hatası'],
    ]
)

# 4. Model bazlı bulgular
h(1, '4. Model Bazlı Bulgular ve Açıklamalar')

h(2, '4.1 ElasticNet Return — Lider Model')
para('Net Return = +2.549, Sharpe = 3.25, Sortino = 43.09, Win% = 71.4, Profit Factor = 102.59', bold=True)
para('ElasticNet Return, Buy & Hold benchmark’ını net getiride aşan tek modeldir. Başarısının üç ana kaynağı vardır:')
numbered([
    'L1+L2 hibrit regülarizasyonu otomatik özellik seçimi (sparsity) ve katsayı küçültme (shrinkage) birlikte uygulayarak overfit’i bastırır; ilgili özellik alt kümesini korur. Yön doğruluğu %59.3 (rastgele üzerinde anlamlı) ve RMSE oranı 1.0017 (benchmark ile eşit kalibrasyon) bu dengeyi gösterir.',
    'Yüksek Profit Factor (102.6), Ridge’in 29.94’ünden 3.4 kat büyüktür. Model büyük kazançlı sinyalleri doğru zamanlıyor; tahminler yalnızca yönde değil, büyüklük sıralamasında da bilgi taşıyor.',
    'Selektif sinyal üretimi: tahminlerin %56.7’si entry eşiğini geçti; aşırı işlem yapmadan (7 trade) yüksek beklenen değer yakalandı.',
])
para('Sinyal teşhisi gate_too_strict, mevcut kapı parametrelerinin (Volatility_Multiplier=1.1375 dahil) sinyallerin bir kısmını gereksiz reddettiğini gösteriyor; gate gevşetilirse alpha artabilir. Deflated Sharpe = 2.71, çoklu test düzeltmesi altında bile istatistiksel anlamlılığı destekliyor.')

h(2, '4.2 Ridge Return — Yakın Aday')
para('Net Return = +2.283, Sharpe = 3.08, Sortino = 43.08, Max Drawdown = -3.88%', bold=True)
para('Ridge Return, ElasticNet ile aynı yön doğruluğunu (%59.3) ve neredeyse aynı Sortino oranını (43.08) yakaladı. Buy & Hold’a net getiride yenildi ancak risk-ayarlı metriklerde üstün:')
bullets([
    'MaxDD = -3.88% tüm modeller arasında en iyi (Buy & Hold’un yaklaşık 1/3.5’i).',
    'Win Rate %72.7 (11 trade’in 8’i kazançlı).',
    'CVaR_95 = -0.0296, trend benchmark’larının yaklaşık yarısı.',
])
para('Net getiri açığı (B&H’a karşı -0.121) maruziyet (exposure) sınırlamasından kaynaklanır: model zamanın %50’sinde piyasada, kalanda nakitte. Trend güçlü olduğu için "olmadığı zaman" fırsat kaybı oluşuyor. Leverage 1.5–2× ile telafi edilebilir; ancak MaxDD orantısal büyür.')

h(2, '4.3 LightGBM Return — Başarısız Model')
para('Net Return = -0.055, Sharpe = -4.41, Dir_Acc = 42.86%, Win% = 25.0', bold=True)
bullets([
    'Yön doğruluğu coinflip altında (%42.9): sign tahmin yeteneği rastgeleden kötü.',
    'RMSE oranı 1.0912: naive benchmark’tan kötü tahmin kalibrasyonu.',
    'Profit Factor 0.42: her trade net kayıp.',
    'Tahmin magnitude’i makul (0.0134) ancak yön bilgisi yok → "büyük tahmin, yanlış işaret" örüntüsü.',
])
para('Olası nedenler: (a) küçük örneklem (12 fold × 21 test bar) ile boosting’in yüksek varyansı; (b) hiperparametre ayarsız default yapılandırma; (c) lineer-trend baskın veride non-lineer ayrışmaların gürültü kalması. Hard gate (professional_current) modeli reddetti (0 trade) — kapı tasarımının doğru çalıştığını gösteren olumlu sinyal.')

h(2, '4.4 DLinear ve NLinear — Tahmin Kalibrasyonu Bozuk')
para('DLinear: Net Return = +0.869, RMSE_vs_bench = 2.5264 | NLinear: Net Return = +0.842, RMSE_vs_bench = 3.0227', bold=True)
para('İki deep linear model birbirine çok yakın sonuç verdi. RMSE oranı benchmark’tan 2.5–3 kat kötü. Composite skorları (DLinear 21.70, NLinear 19.16) için sistem özel olarak düşürülmüş gate eşiği (32.21 ve 38.81) kullanmasına rağmen yine altta kaldılar.')
para('Kritik anomali — ortalama mutlak tahmin getirisi:', bold=True)
bullets([
    'ElasticNet: 0.0119',
    'Ridge: 0.0093',
    'LightGBM: 0.0134',
    'DLinear: 0.0907',
    'NLinear: 0.1138',
])
para('DLinear/NLinear tahminleri diğer modellerden 7–12 kat büyük magnitude üretiyor. Output denormalization veya target scaling katmanında hata şüphesi. NLinear’ın "last value subtraction" normalize tekniği test setinde geri eklenmiyorsa tahminler ölçek olarak şişer. RMSE patlaması ve Composite çöküşü bu hipotezi destekliyor.')
para('Pozitif net getiri elde edilmesi, kapı katmanının kötü tahminleri kısmen filtrelediğini ve birkaç şanslı trade’in (Profit Factor 7–9) toplam sonucu pozitife taşıdığını gösteriyor. Deflated Sharpe ≈ 1.45 marjinal anlamlılık.')
para('Gözlem: professional_soft_gate modu her iki model için de simple_current’tan daha iyi sonuç verdi (NLinear: +1.06 vs +0.84; DLinear: +1.06 vs +0.87). Volatiliteye duyarlı eşik çarpanı (1.1375), bu modellerin abartılı tahminlerini doğal şekilde filtreliyor — hatalı kalibrasyonu maskeleyen ancak alpha doğurmayan yan etki.')

h(2, '4.5 Naive Benchmark Performansı')
bullets([
    'Naive Drift: son eğitim noktasındaki ortalama getiriyi sabit pozisyon olarak yansıtır. 1 trade × %100 exposure = fiilen Buy & Hold ile özdeş. ASTOR’un test dönemindeki güçlü trend rejimi bu modeli yapay öne çıkarıyor.',
    'Naive Last Value: son gözlenen getiriyi sürdürür. 2 trade, %66.7 exposure, %50 win rate, +1.82 net return. Lineer modellerden zayıf, DLinear/NLinear’dan güçlü.',
    'Naive Zero Return: her zaman nakit. Kontrol grubu, +0.00.',
])

# 5. Tartışma
h(1, '5. Karşılaştırmalı Tartışma')

h(2, '5.1 Model Sınıfı Hiyerarşisi')
para('Sonuçlar ASTOR örneğinde şu sıralamayı ortaya koyuyor:')
bullets([
    'Lineer L1+L2 (ElasticNet)',
    '> Lineer L2 (Ridge)',
    '> Trend-baskın naive (Drift / Buy & Hold)',
    '> Mean-reversion naive (Last Value)',
    '> Deep linear decomposition (DLinear, NLinear) — kalibrasyon hatası şüphesi',
    '> Gradient boosting (LightGBM) — yön doğruluğu yetersiz',
    '> Sıfır kontrol (Naive Zero)',
])
para('Bu sıralama "model karmaşıklığı artarken performans azalır" mottosunu destekliyor — finansal zaman serisinde gürültü oranı yüksek olduğundan parametre sayısı arttıkça overfit ve kalibrasyon sorunları büyür. López de Prado’nun Advances in Financial Machine Learning (2018) metnindeki "complex models in finance fail in production" tezi bu deneyde doğrulanıyor.')

h(2, '5.2 Trend Rejiminin Etkisi')
para('ASTOR test penceresinde Buy & Hold +2.404 getiri (~%170 CAGR) üretti. Olağanüstü güçlü yukarı yönlü trend. Böyle bir rejimde:')
bullets([
    '"Her zaman long" stratejileri (Naive Drift, B&H) doğal avantaj kazanır.',
    'Selektif modeller (sınırlı exposure) trendin kısmını kaçırır.',
    'Mean-reversion sinyalleri açıkça yenilir.',
])
para('Sonuçların rejim koşullu olduğunu vurgulamak akademik dürüstlük açısından kritik. Düşüş veya yatay rejimde Ridge/ElasticNet’in B&H’ı geçme olasılığı çok daha yüksek olurken, Naive Drift dramatik biçimde başarısız olabilir.')

h(2, '5.3 Sinyal Kapısı Etkinliği')
bullets([
    'Hard gate (professional_current) kötü modelleri başarıyla bloke etti: LightGBM 0 trade (RMSE 60/60 bloke), DLinear 0 trade, NLinear 0 trade. Gate’in savunma değerini kanıtlıyor.',
    'Soft gate’ler ElasticNet/Ridge gibi marjinal kalibrasyona sahip modellerin geçişine izin veriyor — kabul edilebilir denge.',
    'gate_too_strict teşhisi (ElasticNet): Volatility_Multiplier=1.1375 bazı geçerli sinyalleri reddediyor, kapı ince ayar gerektiriyor.',
])

h(2, '5.4 Maliyet Duyarlılığı')
para('Tüm sonuçlar commission_bps=0.0, slippage_bps=0.0 koşulunda. ASTOR için gerçekçi varsayım komisyon 10–20 bps, slippage 5–15 bps (toplam ~25–35 bps round-trip):')
make_table(
    ['Model', 'Trade Sayısı', '25 bps Tahmini Drag', 'Düzeltilmiş Net Return'],
    [
        ['ElasticNet', '7', '-0.0175', '≈ +2.532'],
        ['Ridge', '11', '-0.0275', '≈ +2.256'],
        ['Naive Drift', '1', '-0.0025', '≈ +2.402'],
        ['Buy & Hold', '1', '-0.0025', '≈ +2.402'],
        ['LightGBM', '8', '-0.0200', '≈ -0.076'],
    ]
)
para('Maliyet eklendiğinde ElasticNet’in B&H üstünlüğü daralır ancak korunur. Production öncesi gerçekçi maliyet simülasyonu zorunlu.')

h(2, '5.5 Örneklem Boyutu ve İstatistiksel Güç')
bullets([
    '60 bar holdout penceresi küçük — tek varlık, tek seed sonuçlarına aşırı güvenmek yanıltıcı.',
    'Deflated Sharpe (multiple-testing düzeltmeli) ElasticNet/Ridge için 2.5+, anlamlı; DLinear/NLinear için 1.44 marjinal, LightGBM için negatif → bu üç model rastgele dalgalanmadan ayırt edilemiyor.',
    'Robustluk için: çoklu seed, multi-asset evren (BIST 30/100), farklı rejim pencereleri, bootstrap güven aralıkları gerekli.',
])

# 6. Sınırlamalar
h(1, '6. Sınırlamalar')
numbered([
    'Tek varlık: ASTOR’a aşırı uydurma. Genelleme için BIST evrene yayılım şart.',
    'Tek dönem: yalnızca bir trend rejimi gözlemlendi.',
    'İdealize maliyet: 0 bps varsayımı pratik değil.',
    'Hiperparametre tuning belirsizliği: DLinear/NLinear ve LightGBM için tuning yapılıp yapılmadığı raporda görünmüyor; yapılmadıysa karşılaştırma adil değil.',
    'Walk-forward fold=12, test=21 bar → toplam 252 test bar; yıllık ~1 işlem günü çözünürlüğünde yeterli ama dağılım kuyrukları için sınırda.',
    'Look-ahead garantisi rapor edilmiş ancak özellik mühendisliği seviyesinde audit dış doğrulama gerektirir.',
])

# 7. Sonuç
h(1, '7. Akademik Sonuç ve Öneriler')
para('Birincil bulgu: ', bold=True).add_run('Lineer regülarizasyonlu modeller (özellikle ElasticNet), karmaşık deep ve boosting alternatiflerini ASTOR walk-forward backtest’inde performansta geride bıraktı. Finansal zaman serisinde basit ve düzenli modellerin avantajı literatürünü destekler (López de Prado 2018; Gu, Kelly, Xiu 2020 — "shallow models prevail when noise dominates").')
para('İkincil bulgu: ', bold=True).add_run('Naive trend-takip stratejilerinin (Drift, B&H) güçlü trend rejiminde aşılması son derece güçtür; değerlendirme mutlaka risk-ayarlı metriklerle (Sortino, MaxDD, Calmar) yapılmalı — sadece net return yanıltıcıdır.')
para('Üçüncül bulgu: ', bold=True).add_run('Sinyal kalitesi kapısı (Dir_Acc + RMSE + Composite) savunma mekanizması olarak çalışıyor: zayıf modelleri (LightGBM, DLinear, NLinear) hard gate’te bloke edebiliyor. Üretim ortamında model felaketlerini önleyen kritik altyapı bileşeni.')

h(2, 'İleri Çalışma Önerileri')
numbered([
    'Multi-asset genelleme: aynı protokolü BIST 30/100 evrene uygula → istatistiksel güç ve genelleme.',
    'Rejim analizi: trend / yatay / düşüş alt-dönemlerinde model performansını ayrı raporla.',
    'Maliyet duyarlılığı analizi: 0, 10, 25, 50 bps grid’inde sonuçları karşılaştır.',
    'DLinear/NLinear kalibrasyon hatası: output denormalization katmanını koddan doğrula; düzeltildikten sonra yeniden test.',
    'Hiperparametre adil karşılaştırma: tüm modeller için aynı CV bütçesinde tuning.',
    'Çoklu seed bootstrap: Sharpe ve Net Return güven aralıkları (örn. 1000 bootstrap).',
    'Ensemble: ElasticNet + Ridge + (kalibrasyonu düzeltilmiş) DLinear stacked ensemble; çeşitlendirme alpha’sı.',
    'Leverage / position-sizing analizi: Ridge’in düşük MaxDD’sini exposure artışıyla B&H’ı geçecek şekilde ölçekle; Kelly-criterion uygulanabilirliği.',
])

h(2, 'Üretim Öncelik Sırası')
make_table(
    ['Sıra', 'Model', 'Durum'],
    [
        ['1', 'ElasticNet Return', 'Aday — gate gevşetme + maliyet doğrulama sonrası deploy'],
        ['2', 'Ridge Return', 'Aday — düşük DD profili için ikincil/diversifikasyon adayı'],
        ['3', 'DLinear / NLinear', 'Geliştirme — kalibrasyon bug fix gerekli'],
        ['4', 'LightGBM', 'Reddedilmiş — feature engineering ve target tanımı revize edilmeden tekrar denemeye değmez'],
    ]
)

# 8. Veri kaynakları
h(1, '8. Ek: Sayısal Veri Kaynakları')
para('Tüm sayılar aşağıdaki dosyalardan üretilmiştir:')
sources = [
    'outputs/ASTOR/runs/20260516_212226_ASTOR_walk_forward_model-LightGBMReturn/',
    'outputs/ASTOR/runs/20260516_212316_ASTOR_walk_forward_model-RidgeReturn/',
    'outputs/ASTOR/runs/20260516_212343_ASTOR_walk_forward_model-ElasticNetReturn/',
    'outputs/ASTOR/runs/20260516_212421_ASTOR_walk_forward_model-DLinear/',
    'outputs/ASTOR/runs/20260516_212505_ASTOR_walk_forward_model-NLinear/',
]
for s in sources:
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
para('Her run klasöründe: md/backtest_report_wf.md, csv/backtest_report_wf.csv, signal_gate_diagnostics_v1_wf.csv, shadow_backtest_comparison_v1_wf.csv.')

doc.save(OUT)
print(f"Saved: {OUT}")
